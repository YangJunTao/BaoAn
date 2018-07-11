# -*- coding: utf-8 -*- 
import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
import docx 
import tkinter.filedialog
from tkinter import *
root = Tk()
root.title("属性数据质量检查")
root.geometry('240x130')
label = Label(root, width=15, height=3, text="选择属性数据").pack()
path = tkinter.filedialog.askopenfilename(filetypes=[("xlsx格式","xlsx")])  


def buttonMessage():
    data_evaluation(path)
    tkinter.messagebox.showinfo(title="完成!",message='已生成属性检查报告至源文件夹下')  
    
    
var = StringVar()
e = Entry(root, textvariable = var)
e.pack()
b = Button(root, text="确定",command = buttonMessage)
var.set(path)
b.pack()
    

def data_evaluation(data_path):
    #创建内存中的word文档对象
    word_file =docx.Document()
    #文件名
    file_name = data_path.split("\\")[-1][:-5]
    #读取
    data = pd.read_excel(data_path)
    print("文件读取正常")
    word_file.add_paragraph(u"文件读取正常")
    
    #列操作
    for col in data.columns.values.tolist():#删除不需要的属性字段
        if col in ['GPSX', 'GPSY',u'经度', u'纬度', 'RecordID', 'MapID']:
            del data[col]
    colList = [column for column in data]#获取列名
    print('属性字段为:')
    word_file.add_paragraph(u"属性字段为:")
    for i in colList:
        print(i);
        word_file.add_paragraph(i)
    
    row = len(data.index)#行数量
    col = len(data.columns)#列数量
    data_null = int(sum(data.isnull().sum(axis=0)))#缺失数据
    
    print('数据量统计如下：')
    print('数据表共有%s行、%s列'%(row, col))
    print('缺失数据共有%s条\r\n'%data_null)
    
    word_file.add_paragraph(u'数据量统计如下：')
    word_file.add_paragraph(u'有%s行、%s列，共计%s条数据'%(row, col, row*col))
    word_file.add_paragraph(u'缺失数据共有%s条\r\n'%data_null)
    '''
    修改列名
    data.rename(columns = {'A':'a', 'B':'b'},inplace = True)
    '''
    
    #画饼状图    
    plt.style.use('ggplot')
    fig = plt.figure(figsize=(8,8))
    labels = [u'所有数据',u'缺失数据']
    sizes = [row*col, data_null ]
    colors=[ '#9999ff', '#ff9999' ]
    explode = [0, 0.1]
    plt.rcParams['font.sans-serif']=['SimHei'] #设置全局字体
    plt.axes(aspect= 'equal')
    plt.xlim( 0, 4)
    plt.ylim( 0, 4)
    plt.pie(x=sizes, labels=labels,textprops = {'fontsize':12, 'color':'k'},
            labeldistance = 1.1,autopct= '%.1f%%',shadow = True,explode=explode, 
            startangle = 90)
    plt.xticks(())
    plt.yticks(())
    plt.title(u"数据缺失率", fontsize = 12)  
    plt.savefig("%s.jpg"%file_name)
    word_file.add_picture("%s.jpg"%file_name)
    plt.show()
    
    
    #字段唯一性检查
    print("字段唯一性检查：重复属性字段有%s个\r\n"%str(len(colList) - len(set(colList))))
    word_file.add_paragraph(u"字段唯一性检查：重复属性字段有%s个\r\n"%str(len(colList) - len(set(colList))))
    
    #数据项唯一性检查
    data_duplicates = np.sum(data.duplicated()==True) #重复数据量
    duplicate_percent = "%.0f%%" % (float(data_duplicates)/float(row)*100)
    print("数据项唯一性检查：重复数据有%s条"%data_duplicates +"，占总数据量的%s\r\n"%duplicate_percent)
    word_file.add_paragraph(u"数据项唯一性检查：重复数据有%s条"%data_duplicates +u"，占总数据量的%s\r\n"%duplicate_percent)
                            
    #字段异常缺失率检查
    print("字段异常缺失率检查：")
    word_file.add_paragraph(u"字段异常缺失率检查：")
    for i in colList:
        col_null = data[i].isnull().sum()
        if col_null > 0.3*row:#筛选缺失率大于30%的字段
            print(str(i)+"字段的缺失率为"+"%.0f%%"%(float(col_null)/float(row)*100))
            word_file.add_paragraph(i+u"字段的缺失率为"+"%.0f%%"%(float(col_null)/float(row)*100))
    print("其余字段的缺失率均低于30%\r\n")
    word_file.add_paragraph(u"其余字段的缺失率均低于30%\r\n")
    
    
    #保存质量报告word文件
    word_file.save(var.get()[:-5]+'.docx')

root.mainloop()         #进入消息循环